#!/usr/bin/env python
"""
Script to create a comprehensive Word document listing all implemented features
and functionalities of the RADOKI IMS project.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_table_header_shading(cell, color):
    """Add background shading to table header cells."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_features_document():
    """Create and populate the features documentation Word document."""
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('RADOKI IMS\nImplemented Features & Functionalities')
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(25, 118, 210)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with date
    subtitle = doc.add_paragraph(f'Documentation Date: {datetime.now().strftime("%B %d, %Y")}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(11)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # Spacing
    
    # Introduction
    intro = doc.add_paragraph(
        'This document provides a comprehensive inventory of all implemented features and '
        'functionalities in the RADOKI Instructional Management System. Each feature is documented '
        'with its description, key functionalities, current status, and testing approach.\n'
    )
    intro.runs[0].font.size = Pt(11)
    
    # Table of contents as text
    doc.add_heading('Contents', level=2)
    contents = [
        'Core & Authentication (4 features)',
        'Academic Features (4 features)',
        'Progress & Certification (3 features)',
        'Promotion & Monetization (3 features)',
        'Communication & Notifications (5 features)',
        'Email & Communication Automation (4 features)',
        'Admin & System Management (2 features)'
    ]
    for item in contents:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Define all features
    features = [
        # Category 1: Core & Authentication
        {
            'category': 'CORE & AUTHENTICATION',
            'features': [
                {
                    'num': 1,
                    'name': 'Authentication & User Roles',
                    'description': 'Complete user authentication system with role-based access control (Student, Instructor, Admin). Users can register, login, and manage their profiles.',
                    'functionalities': [
                        'User registration with email verification',
                        'Role-based permissions (Student/Instructor/Admin)',
                        'Password reset functionality',
                        'Profile customization (bio, photo, qualifications)',
                        'User information management (age, sex, phone, region)',
                        'Session management and logout'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Register new user → Verify role assignment → Login with credentials → Update profile → Verify permissions on dashboard'
                },
                {
                    'num': 2,
                    'name': 'Course Management & Content',
                    'description': 'Comprehensive system for instructors to create, manage, and organize courses with curriculum, pricing, and delivery modes.',
                    'functionalities': [
                        'Course creation with title, description, and curriculum',
                        'Set course pricing and duration',
                        'Choose delivery mode (Online, Offline, Hybrid)',
                        'Organize content with modules and lessons',
                        'Upload and manage course resources',
                        'Set payment deadlines and course dates',
                        'Publish/unpublish courses'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Login as instructor → Create course → Add modules → Upload resources → Set pricing → Publish → Student enrolls'
                },
                {
                    'num': 3,
                    'name': 'Student Dashboard',
                    'description': 'Central hub for students to view enrollments, track progress, access quizzes, and manage their learning activities.',
                    'functionalities': [
                        'View all enrolled courses with approval status',
                        'Monitor course progress with completion percentages',
                        'Access upcoming quizzes and live sessions',
                        'View payment status and receipts',
                        'Check assignment status and grades',
                        'Track attendance records',
                        'Quick links to academic resources'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Login as student → Verify course list → Check progress indicators → Click on course → See resources'
                },
                {
                    'num': 4,
                    'name': 'Instructor Dashboard',
                    'description': 'Management interface for instructors to oversee students, monitor progress, and manage course activities.',
                    'functionalities': [
                        'View all students in courses with completion progress',
                        'Accept/reject student enrollments',
                        'Monitor assignment submissions and grades',
                        'Track quiz attempt results',
                        'Review student attendance records',
                        'Mark courses as complete',
                        'Access student progress analytics'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Login as instructor → View My Students table → Check completion bars → Click student → See analytics'
                }
            ]
        },
        # Category 2: Academic Features
        {
            'category': 'ACADEMIC FEATURES',
            'features': [
                {
                    'num': 5,
                    'name': 'Attendance Tracking System',
                    'description': 'Mobile-responsive attendance system for tracking student attendance at sessions with statistics and reporting. Features responsive design for all screen sizes.',
                    'functionalities': [
                        'Schedule sessions with venue and times',
                        'Mark attendance as present/absent/unrecorded',
                        'View attendance percentage and statistics',
                        'Search and filter attendance records',
                        'Mobile-responsive design (desktop/tablet/mobile)',
                        'Export attendance reports',
                        'Automatic calculation of attendance percentage'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Navigate to attendance → Create session → Mark attendance → Check percentage → Test on mobile device'
                },
                {
                    'num': 6,
                    'name': 'Quizzes Management',
                    'description': 'Full-featured quiz system allowing instructors to create various question types, set pass marks, and track student attempts with detailed scoring.',
                    'functionalities': [
                        'Create quizzes with pass marks and time limits',
                        'Support multiple question types (Multiple Choice, True/False, Short Answer)',
                        'Set max attempts and publishing status',
                        'Track student attempts with time spent',
                        'Auto-calculate scores and pass/fail status',
                        'Display quiz explanations',
                        'Generate quiz reports'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Create quiz → Add questions → Publish → Login as student → Take quiz → Verify score → Check report'
                },
                {
                    'num': 7,
                    'name': 'Assignments & Grading',
                    'description': 'Complete assignment management system with submission tracking, grading, and feedback capabilities.',
                    'functionalities': [
                        'Create assignments with due dates',
                        'Students submit file attachments',
                        'Track submission status (submitted, reviewed, graded, needs resubmission)',
                        'Instructors provide grades and feedback',
                        'Automatic grade notifications',
                        'View submission history',
                        'Bulk grading capabilities'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Create assignment → Student submits → Instructor grades → Student receives notification → Check feedback'
                },
                {
                    'num': 8,
                    'name': 'Lesson Tracking & Progress',
                    'description': 'System to track lesson completion, resource downloads, and student learning progress through lessons and modules.',
                    'functionalities': [
                        'Create lessons within modules',
                        'Track lesson completion status',
                        'Record resource downloads',
                        'Calculate progress percentage per module',
                        'Store lesson completion metadata',
                        'Generate progress reports',
                        'Milestone notifications on completion'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Access course → Complete lesson → Mark as complete → Check progress bar → Download resource → Verify logging'
                }
            ]
        },
        # Category 3: Progress & Certification
        {
            'category': 'PROGRESS & CERTIFICATION',
            'features': [
                {
                    'num': 9,
                    'name': 'Completion Percentage Calculation',
                    'description': 'Automatic calculation of student course completion percentage based on four components (lessons, assignments, quizzes, attendance), each weighted at 25%.',
                    'functionalities': [
                        'Auto-calculate from Lessons (25%), Assignments (25%), Quizzes (25%), Attendance (25%)',
                        'Update signals when student progresses',
                        'Cap at 99% until instructor marks complete',
                        'Real-time progress tracking',
                        'Display completion percentage on dashboards',
                        'Color-coded progress indicators',
                        'Export completion data'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Enroll in course → Complete lessons (track %) → Submit assignments → Take quizzes → Attend sessions → Verify total %'
                },
                {
                    'num': 10,
                    'name': 'Completion Certificates',
                    'description': 'Admin-gated certificate generation system. Certificates are generated only when completion is 100%, instructor marks complete, and admin has enabled certificates for the course.',
                    'functionalities': [
                        'Admin enables certificates per course',
                        'Instructor marks course as complete',
                        'Automatic certificate generation when all conditions met',
                        'Professional PDF certificate with course details',
                        'Student downloads certificate',
                        'Certificate tracking in admin',
                        'Support for auto-generation option'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Admin enables cert → Reach 100% completion → Instructor marks complete → Certificate generates → Download and verify'
                },
                {
                    'num': 11,
                    'name': 'Student Progress Tracking & Analytics',
                    'description': 'Comprehensive progress tracking showing students their academic performance with real-time statistics and visual progress indicators.',
                    'functionalities': [
                        'View personal completion percentage',
                        'Track quiz scores and pass rates',
                        'Monitor assignment grades and feedback',
                        'See attendance records and percentages',
                        'View certificate status',
                        'Compare progress across courses',
                        'Access downloadable progress reports'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Enroll in courses → Complete activities → Navigate to My Progress → Verify all stats are accurate and updated'
                }
            ]
        },
        # Category 4: Promotion & Monetization
        {
            'category': 'PROMOTION & MONETIZATION',
            'features': [
                {
                    'num': 12,
                    'name': 'Coupon Management System',
                    'description': 'Instructor and admin-accessible system for creating and managing promotional discount coupons with flexible discount options and course scoping.',
                    'functionalities': [
                        'Create coupons with unique codes',
                        'Choose discount type (Percentage or Fixed Amount)',
                        'Set validity periods (from/until dates)',
                        'Limit maximum uses or allow unlimited',
                        'Scope coupons to specific courses or all courses',
                        'Enable/disable without deletion',
                        'Track usage statistics and conversion',
                        'View coupon performance metrics'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Instructor creates coupon → Enroll student → Enter coupon code → Verify discount applied → Check usage stats'
                },
                {
                    'num': 13,
                    'name': 'Payment Processing & Receipts',
                    'description': 'Complete payment workflow with receipt upload, approval/rejection process, and payment deadline tracking.',
                    'functionalities': [
                        'Upload payment receipt/proof',
                        'Admin reviews and approves/rejects',
                        'Automatic confirmation emails',
                        'Track payment status and dates',
                        'Manage payment deadlines per course',
                        'Re-upload rejected receipts',
                        'Payment history and tracking',
                        'Overdue payment alerts'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Enroll in course → Upload receipt → Admin approves → Get notification → Check payment status in dashboard'
                },
                {
                    'num': 14,
                    'name': 'Referral System',
                    'description': 'Complete affiliate/referral system allowing students to generate unique referral links, track referrals through signup/enrollment/payment, and earn automatic rewards.',
                    'functionalities': [
                        'Generate unique referral code per student',
                        'Share referral link with friends',
                        'Track referral status through 5 states (Pending → Enrolled → Paid)',
                        'Calculate conversion rates',
                        'Auto-generate rewards on successful payment',
                        'Track earned rewards',
                        'Claim rewards as discounts',
                        'View referral analytics and history'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Copy referral code → Share with friend → Friend registers → Enrolls → Pays → Verify reward created → Check dashboard'
                }
            ]
        },
        # Category 5: Communication & Notifications
        {
            'category': 'COMMUNICATION & NOTIFICATIONS',
            'features': [
                {
                    'num': 15,
                    'name': 'In-App Notification System',
                    'description': 'Comprehensive notification system with 25+ notification types, rich metadata support, and automatic sound reminders every 10 minutes for unread notifications.',
                    'functionalities': [
                        'Support for 25+ notification types',
                        'Rich metadata storage (JSON)',
                        'Mark notifications as read/unread',
                        'Automatic sound alerts every 10 minutes',
                        'Browser desktop notifications (optional)',
                        'Notification center with filtering',
                        'Search and sort notifications',
                        'Pagination support'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Receive notification → Check notification center → Mark as read → Unread notification plays sound → Filter by type'
                },
                {
                    'num': 16,
                    'name': 'Quiz Notifications',
                    'description': 'Automatic notifications sent to all enrolled students when instructors publish new quizzes with detailed quiz information.',
                    'functionalities': [
                        'Auto-notify on quiz publish',
                        'Include quiz details in notification',
                        'Track notification delivery',
                        'Rich metadata with quiz info',
                        'Clickable link to quiz',
                        'Notification filtering',
                        'Sound reminders enabled by default'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Create quiz → Publish → Student receives notification → Check title contains quiz name → Click link → Open quiz'
                },
                {
                    'num': 17,
                    'name': 'Live Session Announcements',
                    'description': 'Automatic notifications for scheduled live sessions sent to all enrolled students with meeting details.',
                    'functionalities': [
                        'Auto-notify on session creation',
                        'Include session time and details',
                        'Add meeting link to notification',
                        'Rich metadata with instructor info',
                        'Clickable notification links',
                        'Calendar integration support',
                        'Weekly digest option'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Schedule live session → Student receives notification → Verify time included → Click link → Join session'
                },
                {
                    'num': 18,
                    'name': 'Resource Upload Notifications',
                    'description': 'Automatic notifications triggered when instructors upload new resources to courses, notifying all enrolled students.',
                    'functionalities': [
                        'Auto-notify on resource upload',
                        'Include resource title and file type',
                        'Add resource metadata',
                        'Direct link to course resources',
                        'File size information',
                        'Notification icon indicating file type',
                        'Download reminder notifications'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Upload course resource → Student receives notification → Click link → Resource page opens → Download file'
                },
                {
                    'num': 19,
                    'name': 'Assignment Notifications',
                    'description': 'Automatic notifications for assignment-related events including new assignments, submissions, grading, and feedback.',
                    'functionalities': [
                        'Notify on new assignment creation',
                        'Notify on submission received',
                        'Notify on assignment graded',
                        'Include due date in notification',
                        'Add feedback snippets',
                        'Grade display in notification',
                        'Resubmission request notifications'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Create assignment → Students notified → Submit → Instructor notified → Grade assignment → Student gets notification with grade'
                }
            ]
        },
        # Category 6: Email & Communication Automation
        {
            'category': 'EMAIL & COMMUNICATION AUTOMATION',
            'features': [
                {
                    'num': 20,
                    'name': 'Payment Approval Email Automation',
                    'description': 'Automatic email sending when admin approves a student payment receipt, notifying them of enrollment approval with course details.',
                    'functionalities': [
                        'Auto-send on payment approval (False → True)',
                        'HTML and plain text email templates',
                        'Include course name and approval message',
                        'Add direct link to course',
                        'Send via configured SMTP',
                        'Error logging and retry handling',
                        'Bulk approval email sending',
                        'Customizable email templates'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Student uploads payment → Admin approves → Check email inbox → Verify contains course name and link → Click link opens course'
                },
                {
                    'num': 21,
                    'name': 'Payment Rejection Email Automation',
                    'description': 'Automatic email notification when admin rejects a payment with customizable rejection reason to guide students on resubmission.',
                    'functionalities': [
                        'Auto-send on payment rejection',
                        'Include customizable rejection reason',
                        'Detailed rejection message template',
                        'Guide students on resubmission',
                        'Direct link to re-upload receipt',
                        'HTML formatted email',
                        'Email status logging',
                        'Bulk rejection notifications'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Student uploads payment → Admin rejects with reason → Check email → Verify reason shown → Click link to resubmit → Upload new receipt'
                },
                {
                    'num': 22,
                    'name': 'Payment Deadline Reminder Email',
                    'description': 'Scheduled management command that automatically sends reminder emails 3 days before payment deadlines to students with pending payments.',
                    'functionalities': [
                        'Send reminders 3 days before deadline',
                        'Target students with no/pending payments',
                        'Use Django management command',
                        'Include deadline date and days remaining',
                        'Add course price and instructor info',
                        'Support dry-run mode for testing',
                        'HTML email templates with styling',
                        'Schedule with cron or Task Scheduler'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Set course with 3-day deadline → Run: python manage.py send_deadline_reminders --dry-run → See what would send → Remove dry-run → Verify emails sent'
                },
                {
                    'num': 23,
                    'name': 'Password Reset Email Automation',
                    'description': 'Automatic email with reset link sent when users request password recovery, supporting both HTML and plain text formats.',
                    'functionalities': [
                        'Send on password reset request',
                        'Include secure reset token link',
                        'HTML and plain text templates',
                        'Token expiration tracking',
                        'One-click reset link',
                        'Customizable email templates',
                        'Django built-in password reset view',
                        'Secure token generation'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Go to login → Forgot password → Enter email → Check inbox → Click reset link → New password form appears → Set new password → Login with new password'
                }
            ]
        },
        # Category 7: Admin & System
        {
            'category': 'ADMIN & SYSTEM MANAGEMENT',
            'features': [
                {
                    'num': 24,
                    'name': 'Admin Activity Logging',
                    'description': 'Comprehensive logging system that tracks all administrator actions including who made changes, what changed, when, and from which IP address.',
                    'functionalities': [
                        'Log all create/update/delete operations',
                        'Record user performing action',
                        'Capture IP address and timestamp',
                        'Track changed fields and values',
                        'Filter logs by action type',
                        'Search logs by user or content',
                        'Export activity reports',
                        'View detailed change history'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Go to admin → Add/edit/delete record → Check Admin Activity Logs → Verify details are recorded correctly'
                },
                {
                    'num': 25,
                    'name': 'System Metrics & Monitoring',
                    'description': 'Automated system metrics collection and monitoring for CPU usage, memory, database activity, and user statistics.',
                    'functionalities': [
                        'Collect CPU and memory usage metrics',
                        'Track database activity statistics',
                        'Monitor active user counts',
                        'Count total courses and enrollments',
                        'Track pending enrollments',
                        'Generate system health reports',
                        'Schedule automated collection',
                        'Export metrics for analysis'
                    ],
                    'status': '✅ Complete',
                    'testing': 'Run: python manage.py collect_system_metrics → Check Core → System Metrics → Verify stats are populated'
                }
            ]
        }
    ]
    
    # Build the document
    for category_group in features:
        doc.add_heading(category_group['category'], level=1)
        
        for idx, feature in enumerate(category_group['features'], 1):
            doc.add_heading(f"Feature {feature['num']}: {feature['name']}", level=2)
            
            # Description
            p = doc.add_paragraph()
            p.add_run('Description: ').bold = True
            p.add_run(feature['description'])
            
            # Key Functionalities
            doc.add_paragraph('Key Functionalities:', style='List Number')
            for func in feature['functionalities']:
                doc.add_paragraph(func, style='List Bullet 2')
            
            # Status
            p = doc.add_paragraph()
            p.add_run('Status: ').bold = True
            status_run = p.add_run(feature['status'])
            if '✅' in feature['status']:
                status_run.font.color.rgb = RGBColor(76, 175, 80)
            
            # Testing Approach
            p = doc.add_paragraph()
            p.add_run('Testing Approach: ').bold = True
            p.add_run(feature['testing'])
            
            doc.add_paragraph()  # Spacing
    
    # Summary page
    doc.add_page_break()
    doc.add_heading('Summary', level=1)
    
    summary_text = [
        'This RADOKI IMS implementation includes 25 key features across 7 major categories:',
        '',
        '• Core & Authentication (4 features): User management, course creation, student and instructor dashboards',
        '• Academic Features (4 features): Attendance tracking, quizzes, assignments, and lesson progress',
        '• Progress & Certification (3 features): Automatic completion calculation, certificates, and progress analytics',
        '• Promotion & Monetization (3 features): Coupons, payment processing, and referral system',
        '• Communication & Notifications (5 features): In-app notifications, quiz/session/resource alerts, and assignments updates',
        '• Email & Communication Automation (4 features): Payment approval/rejection emails, deadline reminders, password reset',
        '• Admin & System (2 features): Activity logging and system metrics collection',
        '',
        'All features are complete and ready for production use. Each feature includes comprehensive functionality '
        'with corresponding testing approaches documented above.'
    ]
    
    for text in summary_text:
        if text:
            doc.add_paragraph(text, style='List Bullet' if text.startswith('•') else None)
        else:
            doc.add_paragraph()
    
    # Save document
    doc_path = r'c:\Users\HP\Documents\Training\radoki_im_system\RADOKI_IMS_Features_and_Functionalities.docx'
    doc.save(doc_path)
    return doc_path

if __name__ == '__main__':
    path = create_features_document()
    print(f'✅ Document created successfully: {path}')
    print(f'Total features documented: 25')
    print(f'Categories: 7')
